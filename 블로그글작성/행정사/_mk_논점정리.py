# -*- coding: utf-8 -*-
"""행정사법 논점정리(업무신고~휴업신고) docx 복원 - 조문은 법령API 현행 원문 기반"""
import json, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SP = os.path.dirname(os.path.abspath(__file__))
OUT = r"c:\Users\USER\OneDrive\바탕 화면\py\tools\블로그글작성\행정사\행정사법_논점정리_업무신고~휴업신고.docx"
JO = json.load(open(os.path.join(SP, "haeng_jo.json"), encoding="utf-8"))

BLUE = RGBColor(0x1E, 0x51, 0xD5)
INK = RGBColor(0x14, 0x1A, 0x26)
SUB = RGBColor(0x5B, 0x65, 0x77)
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia'

def setfont(run, size=10.5, bold=False, color=INK, name="맑은 고딕"):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold; run.font.color.rgb = color
    run._element.rPr.rFonts.set(NS, name)

def para(doc, text="", size=10.5, bold=False, color=INK, indent=0, space_after=4, space_before=0):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.35
    setfont(p.add_run(text), size, bold, color); return p

def nonjo(text):
    t = re.sub(r"\s*<(?:개정|신설|전문개정|본조신설)[^>]*>", "", text)
    return re.sub(r"\s{2,}", " ", t).strip()

def hang(law, jo, n):
    """조문에서 n번째 항 텍스트만 추출"""
    t = nonjo(JO[law + "|" + jo]["text"])
    t = re.sub(r"^제\d+조(?:의\d+)?\([^)]*\)\s*", "", t)
    marks = "①②③④⑤⑥⑦⑧⑨⑩"
    if marks[0] not in t:
        return t if n == 1 else ""
    for p in re.split(r"(?=[①②③④⑤⑥⑦⑧⑨⑩])", t):
        p = p.strip()
        if p.startswith(marks[n-1]):
            return p[1:].strip()
    return ""

doc = Document()
s = doc.sections[0]
s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0); s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

p = para(doc, "행정사법 논점정리", 20, True, BLUE, space_after=2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = para(doc, "업무신고 · 신고확인증 · 사무소 설치·명칭·이전신고 · 폐업·휴업신고", 11, False, SUB, space_after=2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = para(doc, "조문 내용은 국가법령정보센터 현행 법령 원문 기준", 9, False, SUB, space_after=14); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def nonje(no, title, gichul=None):
    para(doc, "[논점 " + str(no) + "] " + title, 15, True, BLUE, space_before=14, space_after=6)
    if gichul: para(doc, "기출 " + gichul, 9.5, False, SUB, indent=0.2, space_after=8)
def h1(t): para(doc, t, 12.5, True, INK, space_before=10, space_after=5)
def h2(t): para(doc, t, 11, True, INK, indent=0.3, space_before=6, space_after=3)
def h3(t): para(doc, t, 10.5, True, SUB, indent=0.6, space_before=3, space_after=2)
def body(t, ind=0.6): para(doc, t, 10.5, False, INK, indent=ind, space_after=3)
def item(t, ind=0.9): para(doc, t, 10.5, False, INK, indent=ind, space_after=2)
def note(t, ind=0.6): para(doc, t, 9.5, False, SUB, indent=ind, space_after=4)

# ── 논점 3
nonje(3, "업무신고",
      "제5회 2문 · 행정사법상 업무신고와 그 수리거부에 관하여 설명하시오.\n"
      "     제8회 2문 · 행정사법상 업무신고의 기준과 행정사업무신고확인증에 관하여 설명하시오.")
h1("Ⅰ. 업무신고(제10조)")
h2("1. 의의(제1항)"); body(hang("행정사법", "제10조", 1))
h2("2. 요건 - 행정사 업무신고 기준(시행령 제20조제1항 각 호)")
body("법 제10조제1항에서 대통령령으로 정하는 행정사 업무신고 기준이란 다음 각 호의 기준을 말한다.")
for t in ["① 법 제6조 각 호의 결격사유에 해당하지 않을 것",
          "② 법 제25조제1항에 따른 실무교육을 이수했을 것",
          "③ 제18조에 따른 행정사 자격증이 있을 것",
          "④ 행정사회에 가입했을 것"]:
    item(t)
h2("3. 절차(시행령 제20조제2항·제3항)")
h3("(1) 제출 대상(제2항 본문)")
body("행정사 업무 신고를 하려는 사람은 행정안전부령으로 정하는 신고서에 다음 각 호의 서류를 첨부하여 "
     "주된 사무소의 소재지를 관할하는 시장등에게 제출해야 한다.", 0.9)
h3("(2) 첨부서류(제2항 각 호)")
item("① 행정사회 회원증 1부", 1.2)
item("② 사진(신청일 전 6개월 이내 촬영, 가로 3센티미터·세로 4센티미터) 1장", 1.2)
h3("(3) 시장등의 확인사항(제3항)")
body("신고서를 제출받은 시장등은 행정사정보시스템을 통하여 다음 각 호의 서류를 확인해야 한다. "
     "다만, 신고하려는 사람이 확인에 동의하지 않는 경우에는 해당 서류의 사본을 첨부하게 해야 한다.", 0.9)
item("① 행정사 자격증 1부", 1.2)
item("② 실무교육 수료증 1부", 1.2)
note("※ 2025. 10. 10. 시행령 개정으로 자격증·실무교육 수료증은 원칙적으로 시스템 확인 사항이 되었고, "
     "신고인이 직접 첨부하는 서류는 회원증과 사진이다.", 0.9)
h2("4. 위반 시 효과(법 제36조제2항제1호)")
body("행정사업무신고 또는 법인업무신고를 하지 아니하고 행정사 업무를 한 자는 1년 이하의 징역 또는 "
     "1천만원 이하의 벌금에 처한다.")

h1("Ⅱ. 수리 거부(제11조)")
h2("1. 수리 거부의 의의(제1항)"); body(hang("행정사법", "제11조", 1))
h2("2. 수리 간주(제2항)"); body(hang("행정사법", "제11조", 2))
h2("3. 이의신청")
h3("(1) 의의(제3항)"); body(hang("행정사법", "제11조", 3), 0.9)
h3("(2) 신고확인증 발급(제4항)"); body(hang("행정사법", "제11조", 4), 0.9)
note("※ 이의신청에 필요한 사항은 행정안전부령으로 정한다(제5항).", 0.9)
doc.add_page_break()

# ── 논점 4
nonje(4, "신고확인증")
h1("Ⅰ. 신고확인증의 발급·재발급(제12조)")
h2("1. 발급(제1항)"); body(hang("행정사법", "제12조", 1))
h2("2. 재발급(제2항)"); body(hang("행정사법", "제12조", 2))
h1("Ⅱ. 신고확인증의 대여 등의 금지(제13조)")
h2("1. 대여 금지(제1항)"); body(hang("행정사법", "제13조", 1))
h2("2. 대여받아 사용 금지(제2항)"); body(hang("행정사법", "제13조", 2))
h2("3. 대여 알선 금지(제3항)"); body(hang("행정사법", "제13조", 3))
h1("Ⅲ. 자격의 취소(제30조제1항제2호)")
body("행정안전부장관은 행정사가 제13조제1항을 위반하여 신고확인증을 양도하거나 대여한 경우에는 "
     "그 자격을 취소하여야 한다.")
note("※ 자격을 취소하려는 경우에는 청문을 하여야 한다(제30조제2항). 청문 필수.")
h1("Ⅳ. 벌칙(제36조제1항제2호)")
body("제13조를 위반하여 신고확인증을 다른 자에게 대여한 행정사·행정사법인과 이를 대여받은 자 또는 "
     "대여를 알선한 자는 3년 이하의 징역 또는 3천만원 이하의 벌금에 처한다.")
doc.add_page_break()

# ── 논점 5
nonje(5, "사무소 설치 · 명칭 · 이전신고")
h1("Ⅰ. 사무소의 설치(제14조)")
h2("1. 행정사사무소(제1항)"); body(hang("행정사법", "제14조", 1))
h2("2. 행정사합동사무소(제2항)")
body("1) 행정사는 업무의 효율적 수행과 공신력을 높이기 위하여 2명 이상의 행정사로 구성된 합동사무소를 "
     "설치할 수 있다.", 0.9)
body("2) 행정사합동사무소를 구성하는 행정사의 수를 넘지 않는 범위에서 주사무소와 분사무소를 설치할 수 있다. "
     "이 경우 주사무소와 분사무소에는 구성행정사가 각각 1명 이상 상근하여야 한다.", 0.9)
note("※ 상근 위반은 업무정지 사유(제32조제1항제2호).", 0.9)
h1("Ⅱ. 사무소의 명칭(제15조)")
h2("1. 명칭 표시 의무(제1항)"); body(hang("행정사법", "제15조", 1))
note("※ 위반 시 100만원 이하 과태료(제38조제2항제2호).")
h2("2. 유사명칭 사용 금지(제2항)"); body(hang("행정사법", "제15조", 2))
note("※ 위반 시 500만원 이하 과태료(제38조제1항제2호).")
h1("Ⅲ. 이전신고(제14조제3항~제5항)")
h2("1. 의의(제3항)"); body(hang("행정사법", "제14조", 3))
h2("2. 신고확인증 발급 및 통지(제4항)"); body(hang("행정사법", "제14조", 4))
h2("3. 처분사유의 승계(제5항)"); body(hang("행정사법", "제14조", 5))
h2("4. 과태료(제38조제2항제1호)")
body("제14조제3항에 따른 사무소 이전신고를 하지 아니한 자에게는 100만원 이하의 과태료를 부과한다.")
doc.add_page_break()

# ── 논점 6
nonje(6, "사무소 폐업 · 휴업신고")
h1("Ⅰ. 폐업신고")
h2("1. 의의(제16조제1항)"); body(hang("행정사법", "제16조", 1))
h2("2. 지위 승계(제33조제1항)"); body(hang("행정사법", "제33조", 1))
h2("3. 행정처분의 승계(제33조제2항)"); body(hang("행정사법", "제33조", 2))
h2("4. 위반사유의 승계(제33조제3항)"); body(hang("행정사법", "제33조", 3))
note("※ 행정처분을 하는 경우에는 폐업한 기간과 폐업의 사유 등을 고려하여 업무정지의 기간을 정하여야 한다(제4항).")
h1("Ⅱ. 휴업신고(제17조)")
h2("1. 의의(제1항)"); body(hang("행정사법", "제17조", 1))
h2("2. 수리(제2항·제3항)")
h3("(1) 수리의 통지(제2항)"); body(hang("행정사법", "제17조", 2), 0.9)
h3("(2) 수리 간주(제3항)"); body(hang("행정사법", "제17조", 3), 0.9)
h2("3. 폐업 간주(제4항)"); body(hang("행정사법", "제17조", 4))
h2("4. 위반 시 효과")
h3("(1) 업무정지 사유(제32조제1항제3호)")
body("제17조제1항에 따른 휴업신고를 하지 아니한 경우, 사무소 소재지를 관할하는 시장등은 6개월의 범위에서 "
     "기간을 정하여 업무의 정지를 명할 수 있다.", 0.9)
note("※ 업무정지처분은 그 사유가 발생한 날부터 3년이 지나면 할 수 없다(제32조제3항).", 1.2)
h3("(2) 법인 소속행정사 등의 요건(제25조의6제3항)")
body("소속행정사 및 법인구성원은 업무정지 중이거나 휴업 중인 사람이 아니어야 한다.", 0.9)

para(doc, "", 9)
p = para(doc, "본 정리는 국가법령정보센터의 현행 조문을 기준으로 작성하였습니다. "
              "시험 대비 시 최신 개정 여부를 확인하시기 바랍니다.", 9, False, SUB)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print("[완료]", OUT)
