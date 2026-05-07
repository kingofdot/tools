#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""terms-review JSON (reviewed/) → 워드(.docx) 변환.

스키마(약관·처리방침 공통):
    {
      "title", "version", "effectiveDate", "preamble",
      "articles": [
        { "no", "title",
          "paragraphs": [
            { "no" | null, "text",
              "label"?,           # 표 제목용
              "items"?: [{no, text}],
              "table"?: {headers, rows}
            }
          ]
        }
      ],
      "appendix"
    }

사용:
    python scripts/json_to_docx.py
    → output/이용약관.docx, output/개인정보처리방침.docx
"""

import json
import sys
from pathlib import Path

# Windows 콘솔(cp949) 에서도 한글·기호가 깨지지 않도록 stdout 을 utf-8 로 재설정
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_KO = "맑은 고딕"
HANG_MARKERS = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"


# ────────────────────────────────────────────────────────
# XML 헬퍼 — 셀 테두리·셀 음영
# ────────────────────────────────────────────────────────
def set_cell_borders(cell, color="888888", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, fill="D9E2F3"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_run_font(run, size_pt=11, bold=False, color=None):
    run.font.name = FONT_KO
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_KO)
    rFonts.set(qn("w:ascii"), FONT_KO)
    rFonts.set(qn("w:hAnsi"), FONT_KO)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def hang_marker(n):
    if isinstance(n, int) and 1 <= n <= 20:
        return HANG_MARKERS[n - 1]
    return f"({n})"


# ────────────────────────────────────────────────────────
# 단락·표 렌더
# ────────────────────────────────────────────────────────
def add_para(doc, text, *, indent_cm=0, bold=False, size=11, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, color=color)
    if italic:
        run.italic = True
    return p


def add_heading(doc, text, *, level=2, size=14):
    """제목 1·2 등을 명명 스타일로 잡고 한국어 폰트도 강제 적용."""
    style_name = "Title" if level == 0 else f"Heading {level}"
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table(doc, headers, rows, *, header_fill="D9E2F3"):
    if not rows:
        return None
    n_cols = max(len(headers), max(len(r) for r in rows))
    headers = list(headers) + [""] * (n_cols - len(headers))
    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        set_run_font(run, size_pt=10, bold=True)
        set_cell_borders(cell)
        shade_cell(cell, fill=header_fill)

    # 데이터
    for r, row in enumerate(rows, start=1):
        for c in range(n_cols):
            val = row[c] if c < len(row) else ""
            cell = table.rows[r].cells[c]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size_pt=10)
            set_cell_borders(cell)
    doc.add_paragraph("")
    return table


def render_paragraph(doc, para):
    """약관·처리방침의 단일 paragraph 항목 렌더."""
    no = para.get("no")
    text = (para.get("text") or "").strip()
    label = (para.get("label") or "").strip()

    # 1) 라벨 (표 제목·섹션명)
    if label:
        add_para(doc, label, bold=True, size=11)
        if text:
            add_para(doc, text)
    elif text:
        prefix = ""
        if isinstance(no, int):
            prefix = f"{hang_marker(no)} "
        add_para(doc, prefix + text)

    # 2) items (호 1. 2. 3. …) + 하위 items (목 가·나·다 …)
    for item in para.get("items") or []:
        item_no = item.get("no")
        item_text = (item.get("text") or "").strip()
        if not item_text:
            continue
        if item_no is None:
            bullet = "• "
        else:
            bullet = f"{item_no}. "
        add_para(doc, bullet + item_text, indent_cm=0.8)
        # 하위 items (목)
        for sub in item.get("items") or []:
            sub_text = (sub.get("text") or "").strip()
            if not sub_text:
                continue
            sub_no = sub.get("no")
            sub_bullet = "" if sub_no is None else f"{sub_no}. "
            add_para(doc, sub_bullet + sub_text, indent_cm=1.6)

    # 3) 표
    table_data = para.get("table")
    if table_data:
        add_table(doc, table_data.get("headers", []), table_data.get("rows", []))


def render_article(doc, article):
    no = article.get("no")
    title = (article.get("title") or "").strip()
    heading = f"제{no}조 ({title})" if title else f"제{no}조"
    add_heading(doc, heading, level=2, size=13)

    for para in article.get("paragraphs") or []:
        render_paragraph(doc, para)


def build_doc(data, output_path):
    doc = Document()

    # ─ 기본 스타일 (Normal) 한국어 폰트 강제 ─
    normal = doc.styles["Normal"]
    normal.font.name = FONT_KO
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_KO)
    rFonts.set(qn("w:ascii"), FONT_KO)
    rFonts.set(qn("w:hAnsi"), FONT_KO)

    # Heading 1·2 한국어 폰트
    for hname in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            h = doc.styles[hname]
            h.font.name = FONT_KO
            rPr = h.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), FONT_KO)
            rFonts.set(qn("w:ascii"), FONT_KO)
            rFonts.set(qn("w:hAnsi"), FONT_KO)
        except KeyError:
            pass

    # ─ 페이지 마진 ─
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─ 제목 ─
    title = data.get("title", "")
    add_heading(doc, title, level=0, size=20)

    # 시행일 / 버전
    eff = data.get("effectiveDate")
    ver = data.get("version")
    sub_parts = []
    if eff:
        sub_parts.append(f"시행일: {eff}")
    if ver:
        sub_parts.append(f"버전: {ver}")
    if sub_parts:
        add_para(doc, " · ".join(sub_parts), size=10,
                 color=RGBColor(0x80, 0x80, 0x80),
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")

    # ─ 전문 ─
    pre = data.get("preamble")
    if pre:
        add_para(doc, pre)
        doc.add_paragraph("")

    # ─ 본문 ─
    for art in data.get("articles") or []:
        render_article(doc, art)

    # ─ 부칙 ─
    appendix = data.get("appendix")
    if appendix:
        doc.add_paragraph("")
        add_heading(doc, "부칙", level=2, size=13)
        add_para(doc, appendix)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(output_path)
        print(f"  ✓ {output_path}")
    except PermissionError:
        # 워드에서 파일이 열려 있을 때 → 타임스탬프 suffix 로 새 파일
        from datetime import datetime
        ts = datetime.now().strftime("%H%M%S")
        alt = output_path.with_name(f"{output_path.stem}__{ts}{output_path.suffix}")
        doc.save(alt)
        print(f"  ⚠ 원본 파일이 열려 있어 새 파일로 저장: {alt}")


def main():
    base = Path(__file__).resolve().parent.parent
    out_dir = base / "output"

    targets = [
        (base / "reviewed" / "service-terms-v2.json", out_dir / "인허가서대리_이용약관.docx"),
        (base / "reviewed" / "privacy-v2.json",       out_dir / "인허가서대리_개인정보처리방침.docx"),
    ]

    for src, dst in targets:
        if not src.exists():
            print(f"  ✗ 입력 파일 없음: {src}")
            continue
        with open(src, encoding="utf-8") as f:
            data = json.load(f)
        print(f"→ {src.name}")
        build_doc(data, dst)


if __name__ == "__main__":
    main()
